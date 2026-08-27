#!/usr/bin/env python3
"""Импортирует первый пул вопросов из DOCX в банк Excel."""

from __future__ import annotations

import hashlib
import re
from collections import Counter
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.worksheet.table import Table, TableStyleInfo

ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "source_pool_001.docx"
TEMPLATE = ROOT / "question_bank_template.xlsx"
OUTPUT = ROOT / "question_bank_pool_001.xlsx"
REPORT = ROOT / "POOL_001_IMPORT_REPORT.md"

LETTERS = ("А", "Б", "В", "Г")
DIFFICULTY = {"Easy": "Базовый", "Medium": "Средний", "Hard": "Повышенный"}
SALES_ROLES = ("РОП / Руководитель отдела продаж", "Менеджеры продаж")
RESPONSIBLE = "РОП / Руководитель отдела продаж"


def topic_for(number: int) -> str:
    if number in {1, 2, 3, 4, 5, 27, 34, 100}:
        return "Компания: история, миссия и культура"
    if number in set(range(6, 29)) | set(range(35, 43)) | set(range(49, 60)) | {101, 102}:
        return "Технологии ЛСТК, ТСП и промышленные здания"
    if number in {44, 45, 46, 47, 48}:
        return "Склады, автосервисы и логистика"
    if number in set(range(29, 34)) | {43, 60, 61, 62, 63, 64, 65, 66} | set(range(89, 99)):
        return "Продажи, CRM и переговоры"
    if number == 67:
        return "Кадровая дисциплина"
    if number == 68:
        return "Маркетинг и содержание публикаций"
    if 69 <= number <= 77:
        return "Проектирование и экспертиза в РК"
    if 78 <= number <= 80:
        return "Договорные модели и лицензирование"
    if 81 <= number <= 88:
        return "Изыскания, приемка и кадастр"
    if number == 99:
        return "Коммерческая тайна и безопасность"
    return "Требует классификации"


def source_for(number: int, topic: str):
    if number == 2:
        return "SRC-REG-003", "Миссия", "Подтверждено по утвержденному РЕГ-003"
    if number == 34:
        return "SRC-REG-001", "Применение Кодекса", "Подтверждено по утвержденному РЕГ-001"
    if number == 99:
        return "SRC-REG-001", "Правило №12 — Безопасность", "Частично подтверждено; точный перечень требует проверки"
    if topic == "Компания: история, миссия и культура":
        return "SRC-ADAPTATION", "не указано", "Источник не предоставлен"
    if topic == "Технологии ЛСТК, ТСП и промышленные здания":
        return "SRC-TECH-LSTK", "не указано", "Технический источник не предоставлен"
    if topic == "Склады, автосервисы и логистика":
        return "SRC-TECH-OBJECTS", "не указано", "Отраслевой источник не предоставлен"
    if topic == "Продажи, CRM и переговоры":
        return "SRC-SALES-BOOK", "не указано", "Книга продаж или регламент не предоставлены"
    if topic == "Кадровая дисциплина":
        return "SRC-HR-LAW", "не указано", "Нормативный источник и редакция не указаны"
    if topic == "Маркетинг и содержание публикаций":
        return "SRC-MARKETING", "не указано", "Регламент маркетинга не предоставлен"
    if topic in {"Проектирование и экспертиза в РК", "Договорные модели и лицензирование", "Изыскания, приемка и кадастр"}:
        return "SRC-LAW-RK", "не указано", "Требуется действующая редакция нормативного источника РК"
    return "SRC-UNKNOWN", "не указано", "Источник не указан"


def extract_questions():
    document = Document(SOURCE)
    answer_rows = document.tables[0].rows[1:]
    answers = {
        int(row.cells[0].text.strip()): {
            "difficulty": row.cells[1].text.strip(),
            "correct": row.cells[2].text.strip(),
            "explanation": row.cells[3].text.strip(),
        }
        for row in answer_rows
    }

    questions = []
    for index, paragraph in enumerate(document.paragraphs):
        match = re.fullmatch(r"Вопрос\s+(\d+)", paragraph.text.strip())
        if not match:
            continue
        number = int(match.group(1))
        wording = document.paragraphs[index + 1].text.strip().replace("What является", "Что является").replace("What представляет", "Что представляет")
        raw_options = document.paragraphs[index + 2].text.strip()
        options = re.findall(r"(?:^|\n)\s*([А-Г])\)\s*(.*?)(?=(?:\n\s*[А-Г]\))|$)", raw_options, re.S)
        if len(options) != 4:
            raise ValueError(f"У вопроса {number} найдено {len(options)} вариантов")
        if number not in answers or answers[number]["correct"] not in LETTERS:
            raise ValueError(f"Нет корректного ключа для вопроса {number}")
        topic = topic_for(number)
        source_code, source_clause, source_note = source_for(number, topic)
        questions.append({
            "number": number,
            "code": f"ВОП-ПУЛ001-{number:03d}",
            "wording": wording,
            "options": options,
            "topic": topic,
            "source_code": source_code,
            "source_clause": source_clause,
            "source_note": source_note,
            **answers[number],
        })
    if len(questions) != 102:
        raise ValueError(f"Ожидалось 102 вопроса, найдено {len(questions)}")
    return questions


def reset_data_rows(ws):
    if ws.max_row > 1:
        ws.delete_rows(2, ws.max_row - 1)


def update_table_ref(ws):
    if ws.tables:
        table = next(iter(ws.tables.values()))
        table.ref = f"A1:{ws.cell(max(ws.max_row, 2), ws.max_column).coordinate}"


def add_report_sheet(wb, questions):
    if "Отчет_пула_001" in wb.sheetnames:
        del wb["Отчет_пула_001"]
    ws = wb.create_sheet("Отчет_пула_001")
    rows = [
        ["Показатель", "Значение"],
        ["Пул", "ПУЛ-001"],
        ["Исходный файл", SOURCE.name],
        ["Контрольная сумма SHA-256", hashlib.sha256(SOURCE.read_bytes()).hexdigest()],
        ["Всего вопросов", len(questions)],
        ["Всего вариантов", sum(len(q["options"]) for q in questions)],
        ["Статус вопросов", "Черновик"],
        ["Назначение по должностям", "; ".join(SALES_ROLES)],
        ["Автор и проверяющий", RESPONSIBLE],
        ["Порог прохождения", "90%"],
        ["Доступность попытки", "1 попытка в месяц"],
        ["Срок действия результата", "по запросу РОПа"],
        ["Хранение источников", "впоследствии в базе данных"],
        ["Прямо подтверждены утвержденными регламентами", 2],
        ["Частично подтверждены утвержденными регламентами", 1],
        ["Требуют предоставления или проверки источника", 99],
        ["Исправлены явные языковые опечатки", "ВОП-ПУЛ001-007, ВОП-ПУЛ001-096"],
    ]
    for difficulty, count in sorted(Counter(q["difficulty"] for q in questions).items()):
        rows.append([f"Сложность: {DIFFICULTY[difficulty]}", count])
    for topic, count in sorted(Counter(q["topic"] for q in questions).items()):
        rows.append([f"Тема: {topic}", count])
    for row in rows:
        ws.append(row)
    for cell in ws[1]:
        cell.fill = PatternFill("solid", fgColor="1F4E78")
        cell.font = Font(color="FFFFFF", bold=True)
    ws.column_dimensions["A"].width = 58
    ws.column_dimensions["B"].width = 90
    ws.freeze_panes = "A2"
    table = Table(displayName="ОтчетПула001", ref=f"A1:B{ws.max_row}")
    table.tableStyleInfo = TableStyleInfo(name="TableStyleMedium2", showRowStripes=True)
    ws.add_table(table)


def build_bank(questions):
    wb = load_workbook(TEMPLATE)
    for sheet in ("Вопросы", "Варианты", "Назначение_по_должности", "Программы", "Источники"):
        reset_data_rows(wb[sheet])

    ws = wb["Вопросы"]
    for q in questions:
        ws.append([
            q["code"], q["topic"], "Один ответ", DIFFICULTY[q["difficulty"]], q["wording"],
            q["explanation"], q["source_code"], q["source_clause"], "Черновик", "0.1",
            RESPONSIBLE, RESPONSIBLE, None,
        ])
    update_table_ref(ws)

    ws = wb["Варианты"]
    for q in questions:
        for order, (letter, text) in enumerate(q["options"], 1):
            ws.append([q["code"], letter, text.strip(), "Да" if letter == q["correct"] else "Нет", 1 if letter == q["correct"] else 0, order])
    update_table_ref(ws)

    ws = wb["Назначение_по_должности"]
    for q in questions:
        for role in SALES_ROLES:
            ws.append([q["code"], role, "Обязательный", "Первый пул для отдела продаж"])
    update_table_ref(ws)

    ws = wb["Программы"]
    ws.append(["ПРОГ-ПУЛ-001", "Комплексное тестирование отдела продаж Star Building", "; ".join(SALES_ROLES), "Все темы пула 001", None, None, 90, 1, "1 месяц", "по запросу РОПа", "Черновик", "0.1"])
    update_table_ref(ws)

    sources = [
        ["SRC-REG-001", "РЕГ-001 Кодекс группы компаний Star Building", "не указано", "Утвержден", "knowledge_base/02_УТВЕРЖДЕНО/full_REG-001_kodeks.md", "2026-08-27"],
        ["SRC-REG-003", "РЕГ-003 Стратегическая основа компании", "не указано", "Утвержден", "knowledge_base/02_УТВЕРЖДЕНО/full_REG-003_strategicheskaya_osnova.md", "2026-08-27"],
        ["SRC-ADAPTATION", "Книга адаптации", "не указано", "Не предоставлен", "не указано", None],
        ["SRC-TECH-LSTK", "Техническая документация ЛСТК и ТСП", "не указано", "Не предоставлен", "не указано", None],
        ["SRC-TECH-OBJECTS", "Технические материалы по типам объектов", "не указано", "Не предоставлен", "не указано", None],
        ["SRC-SALES-BOOK", "Книга продаж и скрипты", "не указано", "Не предоставлен", "не указано", None],
        ["SRC-HR-LAW", "Кадровый регламент или трудовое законодательство", "не указано", "Требует проверки", "не указано", None],
        ["SRC-MARKETING", "Регламент маркетинга", "не указано", "Не предоставлен", "не указано", None],
        ["SRC-LAW-RK", "Нормативно-правовая база строительства Республики Казахстан", "не указано", "Требует действующей редакции", "не указано", None],
        ["SRC-UNKNOWN", "Источник не определен", "не указано", "Не предоставлен", "не указано", None],
    ]
    ws = wb["Источники"]
    for row in sources:
        ws.append(row)
    update_table_ref(ws)

    add_report_sheet(wb, questions)
    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
    wb.save(OUTPUT)


def write_report(questions):
    topic_counts = Counter(q["topic"] for q in questions)
    difficulty_counts = Counter(DIFFICULTY[q["difficulty"]] for q in questions)
    text = [
        "# Отчет импорта первого пула вопросов",
        "",
        f"- Исходник: `{SOURCE.name}`.",
        f"- SHA-256: `{hashlib.sha256(SOURCE.read_bytes()).hexdigest()}`.",
        f"- Вопросов: **{len(questions)}**.",
        f"- Вариантов: **{sum(len(q['options']) for q in questions)}**.",
        "- Все вопросы импортированы со статусом **«Черновик»**.",
        "- Назначение: **отдел продаж** — `РОП / Руководитель отдела продаж` и `Менеджеры продаж`.",
        f"- Автор и проверяющий: **{RESPONSIBLE}**.",
        "- Порог прохождения: **90% правильных ответов**.",
        "- Доступность: **1 попытка в месяц**.",
        "- Срок действия результата: **по запросу РОПа**.",
        "",
        "## Распределение по сложности",
        "",
    ]
    text += [f"- {name}: {count}." for name, count in sorted(difficulty_counts.items())]
    text += ["", "## Распределение по темам", ""]
    text += [f"- {name}: {count}." for name, count in sorted(topic_counts.items())]
    text += [
        "",
        "## Достоверность источников",
        "",
        "- Вопрос 2 прямо подтвержден утвержденным РЕГ-003, раздел «Миссия».",
        "- Вопрос 34 прямо подтвержден утвержденным РЕГ-001, раздел «Применение Кодекса».",
        "- Вопрос 99 частично связан с РЕГ-001, правилом №12, но точный перечень коммерческой тайны требует проверки.",
        "- Остальные 99 вопросов требуют предоставления исходных книг, технической документации либо проверки актуальной редакции нормативных актов.",
        "",
        "## Исправления при нормализации",
        "",
        "- Вопрос 7: `What является` заменено на `Что является`.",
        "- Вопрос 96: `What представляет` заменено на `Что представляет`.",
        "- Смысл, варианты и ключ ответов не изменялись.",
        "",
        "## Вопросы повышенного риска проверки",
        "",
        "- Вопрос 33: финансовое взыскание 30% требует кадровой и правовой проверки.",
        "- Вопрос 45: утверждение об отсутствии специалистов в регионе может быстро устареть.",
        "- Вопросы 69–88: нормы проектирования, экспертизы, лицензирования, приемки и кадастра требуют проверки по действующей редакции законодательства РК.",
        "- Технические абсолютные значения и заявления о долговечности, огнестойкости, температуре, мощности и химической стойкости требуют паспортов, расчетов или утвержденной технической документации.",
        "",
        "## Зафиксированные уточнения",
        "",
        "1. Пул предназначен для отдела продаж: РОП и менеджеров продаж.",
        "2. Ответственный за подготовку и проверку — РОП.",
        "3. Источники будут впоследствии храниться в базе данных; сейчас вопросы остаются черновиками.",
        "4. Порог — 90%; доступна 1 попытка в месяц; результат действует до запроса РОПа на повторное прохождение.",
        "5. Число вопросов и время одной попытки пока не указаны.",
        "",
        "Примечание: РОП указан одновременно автором и проверяющим по решению пользователя. Независимый второй проверяющий пока не назначен.",
    ]
    REPORT.write_text("\n".join(text) + "\n", encoding="utf-8")


def verify(questions):
    wb = load_workbook(OUTPUT, data_only=False)
    assert wb["Вопросы"].max_row == 103
    assert wb["Варианты"].max_row == 409
    assert wb["Назначение_по_должности"].max_row == 205
    assert wb["Программы"].max_row == 2
    assert wb["Источники"].max_row == 11
    question_codes = [row[0].value for row in wb["Вопросы"].iter_rows(min_row=2)]
    assert len(question_codes) == len(set(question_codes)) == 102
    correct_counts = Counter()
    for row in wb["Варианты"].iter_rows(min_row=2, values_only=True):
        if row[3] == "Да":
            correct_counts[row[0]] += 1
    assert len(correct_counts) == 102 and set(correct_counts.values()) == {1}
    assert all(q["source_code"] for q in questions)
    return {
        "questions": 102,
        "options": 408,
        "correct_answers": 102,
        "role_assignments": 204,
        "status": "Черновик",
    }


if __name__ == "__main__":
    pool = extract_questions()
    build_bank(pool)
    write_report(pool)
    print(verify(pool))
    print(OUTPUT)
    print(REPORT)
